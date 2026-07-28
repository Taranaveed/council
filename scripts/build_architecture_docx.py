"""Generate ARCHITECTURE.docx with role → service flowchart layout."""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT / "ARCHITECTURE.docx"

BLUE = "DBEAFE"
GREEN = "D1FAE5"
PURPLE = "E0E7FF"
YELLOW = "FEF3C7"
GRAY = "64748B"


def shade_cell(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, subtitle=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if subtitle:
        p.add_run("\n")
        sub = p.add_run(subtitle)
        sub.font.size = Pt(8)
        sub.font.color.rgb = RGBColor(100, 116, 139)


def merge_row(table, row_idx, col_start, col_end):
    top = table.rows[row_idx].cells[col_start]
    for col in range(col_start + 1, col_end + 1):
        top.merge(table.rows[row_idx].cells[col])


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(f"{bold_prefix} ")
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_info_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, BLUE)
        set_cell_text(cell, h, bold=True, size=9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, val, size=9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()


def add_flowchart(doc):
    cw = [2.2, 2.2, 2.2]
    table = doc.add_table(rows=9, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def full_row(row, text, subtitle, color):
        merge_row(table, row, 0, 2)
        cell = table.rows[row].cells[0]
        shade_cell(cell, color)
        set_cell_text(cell, text, bold=True, subtitle=subtitle)

    def arrow_row(row):
        for col in range(3):
            set_cell_text(table.rows[row].cells[col], "↓", size=12)

    full_row(0, "Synthetic Focus Group App", "Entry point", BLUE)
    arrow_row(1)
    full_row(2, "Login", "Email / password or social sign-in", BLUE)
    arrow_row(3)
    full_row(4, "Role Selection", "User picks who they are", YELLOW)
    arrow_row(5)

    # Role row
    merge_row(table, 6, 0, 1)
    bo = table.rows[6].cells[0]
    shade_cell(bo, GREEN)
    set_cell_text(bo, "Business Owner", bold=True, subtitle="Sellers & product teams")
    buyer = table.rows[6].cells[2]
    shade_cell(buyer, PURPLE)
    set_cell_text(buyer, "Buyer", bold=True, subtitle="Shoppers & end customers")

    arrow_row(7)

    # Services row
    s1 = table.rows[8].cells[0]
    s2 = table.rows[8].cells[1]
    s3 = table.rows[8].cells[2]
    shade_cell(s1, GREEN)
    shade_cell(s2, GREEN)
    shade_cell(s3, PURPLE)
    set_cell_text(s1, "Service 1", bold=True, subtitle="Price Bargaining")
    set_cell_text(s2, "Service 2", bold=True, subtitle="Audience Discovery")
    set_cell_text(s3, "Service 3", bold=True, subtitle="Local Deal Finder")

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "Each service follows the same pipeline: "
        "Form → Live Data → 3-Agent Debate → Judge → Dashboard"
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 116, 139)
    doc.add_paragraph()


def service_section(doc, number, name, role, purpose, inputs, agents, output, api):
    add_heading(doc, f"Service {number} — {name}", level=1)
    add_body(doc, role, bold_prefix="Role:")
    add_body(doc, purpose, bold_prefix="Purpose:")
    add_heading(doc, "User Inputs", level=2)
    add_info_table(doc, ["Field", "Description"], inputs)
    add_heading(doc, "AI Agent Council (3 agents debate before Judge decides)", level=2)
    add_info_table(doc, ["Agent", "What They Focus On"], agents)
    add_heading(doc, "Processing Pipeline", level=2)
    add_info_table(
        doc,
        ["Step", "Action"],
        [
            ["1", "User submits the form"],
            ["2", "Backend fetches live market prices (SerpApi)"],
            ["3", "3 agents debate using real price data (Groq LLM)"],
            ["4", "Judge reads debate and outputs JSON verdict"],
            ["5", "Dashboard shows the final answer to the user"],
        ],
    )
    add_heading(doc, "Final Output on Dashboard", level=2)
    add_body(doc, output)
    add_body(doc, api, bold_prefix="API Route:")
    doc.add_paragraph()


def build_docx():
    doc = Document()

    title = doc.add_heading("Synthetic Focus Group", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_paragraph("System Architecture Map · Version 2.0")
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(100, 116, 139)
    doc.add_paragraph(
        "A multi-agent AI platform where 3 agents debate over live market data "
        "and a Judge delivers the final answer. Users pick a role, choose a service, "
        "and receive actionable results on their dashboard."
    )

    add_heading(doc, "System Flowchart", level=1)
    doc.add_paragraph("Login splits into roles. Each role leads to its own service(s).")
    add_flowchart(doc)

    doc.add_page_break()

    add_heading(doc, "Shared Pipeline (All Services)", level=1)
    add_info_table(
        doc,
        ["Step", "What Happens"],
        [
            ["1 · Login", "User signs in"],
            ["2 · Role Pick", "User chooses Business Owner or Buyer"],
            ["3 · Service Pick", "User selects one of the 3 services below"],
            ["4 · Fill Form", "User enters details specific to that service"],
            ["5 · Live Sync", "Backend fetches real prices before AI runs"],
            ["6 · AI Debate", "3 specialized agents argue over the live data"],
            ["7 · Judge", "Judge agent produces structured JSON verdict"],
            ["8 · Dashboard", "User sees the final recommendation"],
        ],
    )

    doc.add_page_break()

    service_section(
        doc, 1, "Price Bargaining", "Business Owner",
        "Help sellers find the optimal price for their product using live competitor data.",
        [
            ["Product Specs", "Key features, materials, and product details"],
            ["Estimated Price Range", "Minimum and maximum price under consideration"],
        ],
        [
            ["Premium Maximizer", "Argues for higher margins and premium positioning"],
            ["Volume Discounter", "Argues for lower prices to drive sales volume"],
            ["Market Benchmark Proxy", "Uses live competitor prices as neutral ground truth"],
        ],
        "Recommended selling price with confidence score, market comparison, and debate summary.",
        "POST /api/modes/price-bargaining",
    )

    service_section(
        doc, 2, "Audience Discovery", "Business Owner",
        "Help sellers identify who would buy their product and why.",
        [
            ["Product Name", "Name of the product being analyzed"],
            ["Problem It Solves", "Primary problem the product addresses for customers"],
        ],
        [
            ["Demographic Scout", "Identifies age, income, location, and lifestyle segments"],
            ["Psychographic Analyst", "Analyzes values, motivations, and buying triggers"],
            ["Utility Specialist", "Evaluates practical use cases and feature fit"],
        ],
        "Three ideal buyer persona profiles with motivations, triggers, and channel recommendations.",
        "POST /api/modes/audience-discovery",
    )

    service_section(
        doc, 3, "Local Deal Finder", "Buyer",
        "Help shoppers find the best real-world deal based on live local listings.",
        [
            ["Item Name", "Product the buyer wants to purchase"],
            ["Max Budget", "Maximum amount willing to spend (optional)"],
            ["Location", "City or region for localized price search"],
        ],
        [
            ["Thrift Advocate", "Finds the lowest price and best overall value"],
            ["Risk Analyst", "Evaluates vendor trust, warranty, and return policies"],
            ["Contextual Persona", "Matches the deal to buyer budget, location, and urgency"],
        ],
        "Buy or Pass verdict, best deal card (store, price, link), alternatives list, and agent reasoning.",
        "POST /api/modes/deal-finder",
    )

    doc.add_page_break()

    add_heading(doc, "Tech Stack", level=1)
    add_info_table(
        doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Frontend", "React + Vite + TypeScript", "User dashboards and forms"],
            ["Backend API", "FastAPI (Python)", "Request routing and agent orchestration"],
            ["Live Prices", "Node.js + SerpApi", "Fetch real market listings before AI runs"],
            ["AI Engine", "Groq · Llama 3.3 70B", "Agent debates and Judge verdict"],
        ],
    )

    doc.save(str(DOCX_PATH))
    print(f"Created {DOCX_PATH}")


if __name__ == "__main__":
    build_docx()
